from __future__ import annotations

import hashlib
import json
import logging
from pathlib import Path
from typing import Any
from datetime import datetime, timezone
from threading import RLock

# ---- Chroma requires sqlite3 >= 3.35.0 on Linux/WSL.
# If the Python stdlib sqlite3 module is linked against an older libsqlite3,
# use the bundled pysqlite3-binary instead (no system upgrade required).
try:
    import sqlite3  # noqa: F401

    def _sqlite_version_tuple(v: str) -> tuple[int, int, int]:
        parts = (v.split(".") + ["0", "0", "0"])[:3]
        return int(parts[0]), int(parts[1]), int(parts[2])

    _min_sqlite = (3, 35, 0)
    _have_sqlite = _sqlite_version_tuple(sqlite3.sqlite_version)  # type: ignore[attr-defined]
    if _have_sqlite < _min_sqlite:
        import sys

        import pysqlite3  # type: ignore[import-not-found]

        sys.modules["sqlite3"] = pysqlite3
except Exception:
    # If any of this fails, we let the normal import error surface when Chroma loads.
    pass

from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.config import settings
from app.prompts import load_prompt

logger = logging.getLogger(__name__)

DOC_SOURCE = "doc.txt"

# ---- Indexing coordination
# Background indexing can run concurrently with requests in FastAPI. We keep it simple and
# guard Chroma writes + manifest updates with a process-local lock (single-process server).
_INDEX_LOCK = RLock()


def _project_root() -> Path:
    return Path(__file__).resolve().parent.parent


def _data_dir() -> Path:
    return _project_root() / "data"


def _doc_path() -> Path:
    return _data_dir() / DOC_SOURCE


def _raw_dir() -> Path:
    """
    Optional multi-document input directory.

    If it exists and contains supported files, we index those instead of the single hardcoded doc.txt.
    This preserves existing single-doc behavior by default (no ./data/raw directory required).
    """
    return _data_dir() / "raw"


def _supported_raw_files() -> list[Path]:
    """
    Return supported raw document files (txt/md) from ./data/raw.

    Notes:
    - Sorted for deterministic indexing behavior.
    - If ./data/raw does not exist or is empty, we fall back to the legacy single doc.txt.
    """
    raw = _raw_dir()
    if not raw.exists():
        return []
    files: list[Path] = []
    for p in raw.iterdir():
        if not p.is_file():
            continue
        # Supported ingest/index formats (uploads + URL saves):
        # - .txt/.md: direct
        # - .pdf/.docx: extracted during indexing
        if p.suffix.lower() in {".txt", ".md", ".pdf", ".docx"}:
            files.append(p)
    return sorted(files, key=lambda x: x.name.lower())


def _index_sources() -> list[Path]:
    """
    Determine which document sources should be indexed.

    Behavior:
    - If ./data/raw contains any supported docs, index those (multi-doc mode).
    - Otherwise, index the legacy ./data/doc.txt only (single-doc mode).
    """
    raw_files = _supported_raw_files()
    if raw_files:
        return raw_files
    return [_doc_path()]


def _index_dir() -> Path:
    return _data_dir() / "index"


def _manifest_path() -> Path:
    """
    Per-source incremental indexing manifest.

    This tracks the file hash + chunking params so we can skip re-indexing unchanged docs.
    """
    return _index_dir() / "manifest.json"


def _read_manifest() -> dict[str, Any]:
    """
    Read the manifest (or return an empty structure if missing/corrupt).

    Shape:
    {
      "version": 1,
      "sources": {
        "<filename>": { ... entry ... }
      }
    }
    """
    path = _manifest_path()
    if not path.exists():
        return {"version": 1, "sources": {}}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        if not isinstance(data, dict):
            return {"version": 1, "sources": {}}
        if "sources" not in data or not isinstance(data.get("sources"), dict):
            data["sources"] = {}
        data.setdefault("version", 1)
        return data
    except Exception:
        logger.exception("Failed to read manifest; continuing with empty manifest (will re-index as needed).")
        return {"version": 1, "sources": {}}


def _write_manifest(data: dict[str, Any]) -> None:
    """
    Persist the manifest to disk.

    Notes:
    - This file lives under ./data/index and is intentionally *not* committed.
    """
    _index_dir().mkdir(parents=True, exist_ok=True)
    _manifest_path().write_text(json.dumps(data, indent=2), encoding="utf-8")


def _fingerprint_path() -> Path:
    """
    Stores the fingerprint of the currently-indexed doc/settings so we can detect changes
    and rebuild the persistent Chroma index automatically.
    """
    return _index_dir() / "doc_fingerprint.json"


def load_doc() -> str:
    return _doc_path().read_text(encoding="utf-8")


def _compute_fingerprint() -> dict[str, Any]:
    """
    Fingerprint includes document contents + indexing-relevant settings.

    Notes:
    - We hash the raw bytes to avoid any encoding-normalization surprises.
    - If any of these fields change, we should rebuild the index.
    """
    sources = _index_sources()
    sources_fp: list[dict[str, str]] = []
    for p in sources:
        # Hash raw bytes to avoid encoding normalization differences.
        b = p.read_bytes()
        sources_fp.append({"source": p.name, "sha256": hashlib.sha256(b).hexdigest()})
    return {
        "sources": sources_fp,
        "chunk_size": int(settings.chunk_size),
        "chunk_overlap": int(settings.chunk_overlap),
        "embedding_model": str(settings.embedding_model),
    }


def _read_fingerprint() -> dict[str, Any] | None:
    path = _fingerprint_path()
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        logger.exception("Failed to read index fingerprint; forcing re-index.")
        return None


def _write_fingerprint(fp: dict[str, Any]) -> None:
    _fingerprint_path().write_text(json.dumps(fp, indent=2), encoding="utf-8")


def _sha256_bytes(b: bytes) -> str:
    """
    Compute a sha256 hex digest for raw bytes.
    """
    return hashlib.sha256(b).hexdigest()


def _now_iso() -> str:
    """
    UTC timestamp used for manifest/job bookkeeping.
    """
    return datetime.now(timezone.utc).isoformat()


def _extract_text_from_file(path: Path) -> str:
    """
    Extract UTF-8 text from supported document types.

    Supported:
    - .txt/.md: read as utf-8 (errors='ignore' to avoid crashing on odd encodings)
    - .pdf: use pypdf text extraction per page
    - .docx: use python-docx paragraph text
    """
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")

    if suffix == ".pdf":
        # pypdf extraction can be imperfect, but it's lightweight and good enough for RAG.
        from pypdf import PdfReader  # type: ignore[import-not-found]

        reader = PdfReader(str(path))
        parts: list[str] = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                # Keep going even if one page fails.
                parts.append("")
        return "\n\n".join(p for p in parts if p).strip()

    if suffix == ".docx":
        from docx import Document as DocxDocument  # type: ignore[import-not-found]

        doc = DocxDocument(str(path))
        parts = [p.text for p in doc.paragraphs if (p.text or "").strip()]
        return "\n".join(parts).strip()

    raise ValueError(f"Unsupported file type for indexing: {path.name}")


def _split_sections(text: str) -> list[tuple[str, str]]:
    """
    Split a document into (section_title, section_text) using a hierarchical approach.

    Supports:
    - Markdown ATX headings: "# Title", "## Title", ...
    - Setext headings:
        Title
        -----
      (or =====)

    If no headings are detected, returns a single ("", full_text) section.
    """
    lines = (text or "").splitlines()
    sections: list[tuple[str, list[str]]] = []
    cur_title = ""
    cur_lines: list[str] = []

    def _flush() -> None:
        nonlocal cur_title, cur_lines
        if cur_lines:
            sections.append((cur_title.strip(), cur_lines[:]))
        cur_lines = []

    i = 0
    while i < len(lines):
        ln = lines[i]
        stripped = ln.strip()

        # ATX heading: "# Title"
        if stripped.startswith("#"):
            title = stripped.lstrip("#").strip()
            if title:
                _flush()
                cur_title = title
                i += 1
                continue

        # Setext heading:
        # Title
        # ------
        if i + 1 < len(lines):
            next_ln = lines[i + 1].strip()
            if next_ln and set(next_ln) <= {"=", "-"} and len(next_ln) >= 3 and stripped:
                _flush()
                cur_title = stripped
                i += 2
                continue

        cur_lines.append(ln)
        i += 1

    _flush()

    # If we never created a titled section, fall back to one big section.
    if not any(t for t, _ in sections):
        return [("", text)]

    out: list[tuple[str, str]] = []
    for title, body_lines in sections:
        body = "\n".join(body_lines).strip()
        if not body:
            continue
        out.append((title, body))
    return out or [("", text)]


def chunk_doc(text: str, *, source: str) -> list[Document]:
    """
    Chunk a single document into stable chunk IDs and attach metadata.

    Requirements:
    - Chunk IDs are deterministic and stable: chunk-00, chunk-01, ...
    - Metadata includes:
      - source: filename
      - section: section header title (if available)
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", " ", ""],
    )

    docs: list[Document] = []
    chunk_index = 0

    for section_title, section_text in _split_sections(text):
        # Split within section for better locality; include the title as context if present.
        prefix = f"{section_title}\n\n" if section_title else ""
        chunks = splitter.split_text(prefix + section_text)
        for chunk in chunks:
            chunk_id = f"chunk-{chunk_index:02d}"
            chunk_index += 1
            # IMPORTANT: Chroma metadata values must be scalar (str/int/float/bool) and not None.
            # To satisfy the "section is null when unavailable" requirement, we omit the key entirely
            # when there is no section title. Consumers should treat missing as null.
            metadata: dict[str, Any] = {"source": source, "chunk_id": chunk_id}
            if section_title:
                metadata["section"] = section_title
            docs.append(
                Document(
                    page_content=chunk,
                    metadata=metadata,
                )
            )
    return docs


def _get_embeddings() -> OpenAIEmbeddings:
    # langchain-openai reads OPENAI_API_KEY env var; we keep settings for clarity/validation.
    if not settings.openai_api_key:
        logger.warning("OPENAI_API_KEY is empty; OpenAI calls will fail until it is set.")
    return OpenAIEmbeddings(model=settings.embedding_model)


def _get_llm() -> ChatOpenAI:
    if not settings.openai_api_key:
        logger.warning("OPENAI_API_KEY is empty; OpenAI calls will fail until it is set.")
    return ChatOpenAI(model=settings.model, temperature=0)


def get_vectorstore() -> Chroma:
    _index_dir().mkdir(parents=True, exist_ok=True)
    return Chroma(
        collection_name="doc",
        persist_directory=str(_index_dir()),
        embedding_function=_get_embeddings(),
    )


def index_if_needed() -> None:
    """
    Ensure the persistent Chroma index exists and is up to date.

    Behavior (incremental, best-practice):
    - Load the persisted vector store under ./data/index
    - Load the persisted manifest under ./data/index/manifest.json
    - For each source file:
      - If file hash + chunking params unchanged: skip
      - Else: delete old chunks for that source and re-index only that file
    - If a previously-indexed source file was removed: delete its chunks and remove it from the manifest

    Notes:
    - We intentionally do NOT full-rebuild on startup or per request.
    - We keep the legacy doc_fingerprint.json helpers around (older versions), but the manifest is authoritative.
    """
    index_scan_incremental()


def index_path_incremental(
    path: Path,
    *,
    source_type: str,
    _vs: Chroma | None = None,
    _manifest: dict[str, Any] | None = None,
    write_manifest: bool = True,
) -> dict[str, Any]:
    """
    Incrementally index a single source file into the persisted Chroma collection.

    Returns a small dict with {status, skipped, filename, num_chunks}.

    Requirements:
    - If file hash and chunk params unchanged: skip
    - If changed: delete old chunks for that source and re-index only that file
    """
    with _INDEX_LOCK:
        vs = _vs or get_vectorstore()
        manifest = _manifest or _read_manifest()
        # Ensure sources dict exists.
        sources: dict[str, Any] = manifest.setdefault("sources", {})

        st = path.stat()
        filename = path.name
        chunk_size = int(settings.chunk_size)
        chunk_overlap = int(settings.chunk_overlap)

        existing = sources.get(filename) or {}
        # Cheap short-circuit: if mtime+size and chunk params match, trust it and skip hashing.
        if (
            existing
            and float(existing.get("last_modified", -1)) == float(st.st_mtime)
            and int(existing.get("size_bytes", -1)) == int(st.st_size)
            and int(existing.get("chunk_size", -1)) == chunk_size
            and int(existing.get("chunk_overlap", -1)) == chunk_overlap
        ):
            return {"status": "ok", "skipped": True, "filename": filename, "num_chunks": int(existing.get("num_chunks", 0))}

        raw_bytes = path.read_bytes()
        content_hash = _sha256_bytes(raw_bytes)

        # If content hash and chunk params match, skip (but update mtime/size for correctness).
        if (
            existing
            and str(existing.get("content_hash", "")) == content_hash
            and int(existing.get("chunk_size", -1)) == chunk_size
            and int(existing.get("chunk_overlap", -1)) == chunk_overlap
        ):
            existing["last_modified"] = float(st.st_mtime)
            existing["size_bytes"] = int(st.st_size)
            existing["indexed_at"] = existing.get("indexed_at") or _now_iso()
            sources[filename] = existing
            if write_manifest:
                _write_manifest(manifest)
            return {"status": "ok", "skipped": True, "filename": filename, "num_chunks": int(existing.get("num_chunks", 0))}

        # Delete old chunks for this source (if any), then re-index.
        try:
            vs._collection.delete(where={"source": filename})  # noqa: SLF001
        except Exception:
            logger.exception("Failed deleting prior chunks for source=%s; continuing (may duplicate).", filename)

        text = _extract_text_from_file(path)
        docs = chunk_doc(text, source=filename)
        ids = [f"{filename}::{d.metadata['chunk_id']}" for d in docs]
        vs.add_documents(docs, ids=ids)
        try:
            vs.persist()
        except Exception:
            pass

        sources[filename] = {
            "filename": filename,
            "source_type": str(source_type),
            "last_modified": float(st.st_mtime),
            "size_bytes": int(st.st_size),
            "content_hash": content_hash,
            "chunk_size": chunk_size,
            "chunk_overlap": chunk_overlap,
            "num_chunks": len(docs),
            "indexed_at": _now_iso(),
        }
        if write_manifest:
            _write_manifest(manifest)
        return {"status": "ok", "skipped": False, "filename": filename, "num_chunks": len(docs)}


def index_scan_incremental() -> dict[str, Any]:
    """
    Incrementally scan expected sources and index only those missing/changed.

    This is safe to call on startup and on each request; unchanged docs are skipped.
    """
    with _INDEX_LOCK:
        _index_dir().mkdir(parents=True, exist_ok=True)
        vs = get_vectorstore()
        manifest = _read_manifest()
        sources: dict[str, Any] = manifest.setdefault("sources", {})

        desired = {p.name: p for p in _index_sources()}

        # Remove entries whose files no longer exist.
        removed = [name for name in list(sources.keys()) if name not in desired]
        for name in removed:
            try:
                vs._collection.delete(where={"source": name})  # noqa: SLF001
            except Exception:
                logger.exception("Failed deleting chunks for removed source=%s; continuing.", name)
            sources.pop(name, None)

        # Index/update current sources.
        updated: list[dict[str, Any]] = []
        skipped: list[str] = []
        for name, p in desired.items():
            # Preserve known source_type from manifest where available.
            existing = sources.get(name) or {}
            source_type = str(existing.get("source_type") or ("url" if name.startswith("url_") else "upload"))
            res = index_path_incremental(p, source_type=source_type, _vs=vs, _manifest=manifest, write_manifest=False)
            if res.get("skipped"):
                skipped.append(name)
            else:
                updated.append(res)

        _write_manifest(manifest)
        return {"status": "ok", "updated": updated, "skipped": skipped, "removed": removed}


def get_retriever(top_k: int) -> Any:
    vs = get_vectorstore()
    # Prefer MMR (more diverse results) where supported by the vectorstore wrapper.
    # If unsupported, LangChain will fall back to similarity search.
    return vs.as_retriever(search_type="mmr", search_kwargs={"k": top_k})


def _citation_key(doc: Document) -> str:
    """
    Canonical key for (source, chunk_id) pairs used in citations and deduplication.
    """
    src = str(doc.metadata.get("source", DOC_SOURCE))
    cid = str(doc.metadata.get("chunk_id", "chunk-??"))
    return f"{src}#{cid}"


def _select_diverse(
    candidates: list[tuple[Document, float]],
    *,
    top_k: int,
) -> list[tuple[Document, float]]:
    """
    Select up to top_k candidates with basic diversity constraints.

    Notes:
    - We deduplicate by citation key first.
    - We then avoid returning everything from the same section when possible.
    """
    # Deduplicate by citation key first (multi-query merges will produce duplicates).
    seen: set[str] = set()
    deduped: list[tuple[Document, float]] = []
    for d, s in candidates:
        key = _citation_key(d)
        if key in seen:
            continue
        seen.add(key)
        deduped.append((d, s))

    if len(deduped) <= top_k:
        return deduped

    # Soft cap: at most 2 per (source, section) group when possible.
    out: list[tuple[Document, float]] = []
    per_group: dict[tuple[str, str], int] = {}
    for d, s in deduped:
        src = str(d.metadata.get("source", DOC_SOURCE))
        # section may be None; treat as empty for grouping.
        sec = str(d.metadata.get("section") or "")[:120]
        group = (src, sec)
        if per_group.get(group, 0) >= 2 and len(per_group) > 1:
            continue
        out.append((d, s))
        per_group[group] = per_group.get(group, 0) + 1
        if len(out) >= top_k:
            break

    # If we were too strict and didn't fill, backfill from remaining deduped.
    if len(out) < top_k:
        chosen = {_citation_key(d) for d, _ in out}
        for d, s in deduped:
            if _citation_key(d) in chosen:
                continue
            out.append((d, s))
            if len(out) >= top_k:
                break
    return out


def retrieve_with_scores(question: str, top_k: int) -> list[tuple[Document, float]]:
    """
    Retrieve documents with a score field.

    Behavior:
    - Prefer MMR (diverse) where available.
    - Fall back to similarity_search_with_score.
    """
    vs = get_vectorstore()
    fetch_k = max(int(top_k) * 4, 20)
    try:
        docs = vs.max_marginal_relevance_search(question, k=top_k, fetch_k=fetch_k)
        # MMR does not naturally expose a score; return a deterministic placeholder.
        return [(d, 0.0) for d in docs]
    except Exception:
        # Chroma returns "distance" (lower is better) for many distance metrics.
        candidates = vs.similarity_search_with_score(question, k=fetch_k)
        return _select_diverse(candidates, top_k=top_k)


def _build_context(retrieved: list[tuple[Document, float]]) -> str:
    parts: list[str] = []
    for doc, _score in retrieved:
        src = str(doc.metadata.get("source", DOC_SOURCE))
        chunk_id = str(doc.metadata.get("chunk_id", "chunk-??"))
        parts.append(f"[{src}#{chunk_id}]\n{doc.page_content}".strip())
    return "\n\n".join(parts).strip()


def _has_citation_token(text: str) -> bool:
    # Generalized: allow [<filename>#chunk-XX] for multi-doc mode.
    t = text or ""
    return ("[" in t) and ("#chunk-" in t)


def answer_question(
    question: str,
    top_k: int = 4,
    debug: bool = False,
) -> dict[str, Any]:
    index_if_needed()

    retrieved = retrieve_with_scores(question, top_k=top_k)
    citations = []
    debug_retrieved = []

    for doc, score in retrieved:
        chunk_id = str(doc.metadata.get("chunk_id", "chunk-??"))
        src = str(doc.metadata.get("source", DOC_SOURCE))
        snippet = doc.page_content.strip().replace("\n", " ")
        snippet = snippet[:240] + ("..." if len(snippet) > 240 else "")
        citations.append({"source": src, "chunk_id": chunk_id, "snippet": snippet})
        debug_retrieved.append({"chunk_id": chunk_id, "text": doc.page_content, "score": float(score)})

    if len(retrieved) == 0:
        out: dict[str, Any] = {
            "answer": "I can’t find that in the provided documentation.",
            "citations": citations,
        }
        if debug:
            out["debug"] = {"retrieved": debug_retrieved}
        return out

    system = load_prompt("system")
    answer_tmpl = load_prompt("answer")
    context = _build_context(retrieved)

    llm = _get_llm()
    messages = [
        SystemMessage(content=system),
        HumanMessage(
            content=answer_tmpl.format(
                question=question,
                context=context,
            )
        ),
    ]

    resp = llm.invoke(messages)
    answer = (resp.content or "").strip()

    # Deterministic post-check: refuse unless there is at least one citation token AND
    # all cited (source, chunk_id) pairs were actually retrieved.
    if not _has_citation_token(answer):
        answer = "I can’t find that in the provided documentation."
    else:
        import re

        # Extract strict citations: [<filename>#chunk-XX]
        cited_pairs = re.findall(r"\[([^\]#]+)#(chunk-\d+)\]", answer)
        if not cited_pairs:
            answer = "I can’t find that in the provided documentation."
        else:
            allowed = {_citation_key(d) for d, _ in retrieved}
            for src, cid in cited_pairs:
                if f"{src}#{cid}" not in allowed:
                    answer = "I can’t find that in the provided documentation."
                    break

    out2: dict[str, Any] = {"answer": answer, "citations": citations}
    if debug:
        out2["debug"] = {"retrieved": debug_retrieved}
    return out2


